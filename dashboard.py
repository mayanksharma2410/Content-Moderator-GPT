import openai
import streamlit as st
import docx
from io import StringIO 
from PIL import Image
import io
import zipfile

openai.api_key = st.secrets["API_SECRET"]

# Load the Document
def load_doc(uploaded_file):
    doc = docx.Document(uploaded_file)
    # Get the text from the document
    text_list = []
    doc_text = ""
    for para in doc.paragraphs:
        text_list.append(para.text)
        doc_text += para.text + " "

    # Ommiting empty spaces in the text
    text_list = list(filter(bool, text_list))
    return doc_text, text_list[0]

# Save the Document
def save_doc(answer, action):
    # Create a new document
    document = docx.Document()

    # Add a heading to the document
    # document.add_heading(heading, level=0)

    # Add a paragraph of text
    document.add_paragraph(answer)

    return document

def save_multiple_doc(documents):
    # Create a new in-memory buffer for the zip archive
    zip_buffer = io.BytesIO()
    
    # Create a new zip archive and add the Word documents to it
    with zipfile.ZipFile(zip_buffer, mode='w', compression=zipfile.ZIP_DEFLATED) as archive:
        for i, document in enumerate(documents):
            # Save each document to a temporary buffer
            doc_buffer = io.BytesIO()
            document.save(doc_buffer)
            doc_buffer.seek(0)

            # Add the document to the zip archive with a filename based on its index
            archive.writestr(f'document{i+1}.docx', doc_buffer.getbuffer())

    # Reset the buffer's position
    zip_buffer.seek(0)

    # Use Streamlit's download_button to download the zip archive
    st.download_button(
        label='Download documents as zip',
        data=zip_buffer,
        file_name='documents.zip',
        mime='application/zip'
    )

# Function to use ChatGPT api
def chatgpt_api(query):
    output = openai.ChatCompletion.create(model = "gpt-3.5-turbo",
                                                # system, user, assistant
                                                messages = [{"role" : "user", "content" : query}]
                                            )
    answer = output['choices'][0]["message"]["content"]
    return answer

# Function to use ChatGPT to generate thread
def gen_thread(article):
    query = article + " " + "make twitter thread of 5 tweets from this article in the same language"
    output = openai.ChatCompletion.create(model = "gpt-3.5-turbo",
                                                # system, user, assistant
                                                messages = [{"role" : "user", "content" : query}]
                                            )
    answer = output['choices'][0]["message"]["content"]
    return answer

# Text and images
st.title("Content Correction with GPT")
gpt_image = Image.open('gpt-image.jpg')
st.image(gpt_image)
st.subheader("Intoduction")
st.write("This tool leverage the power of GPT-3, Natural Language Processing to process text and translate the document into 3 languages supported, i.e. English, Hindi & Urdu. The tool even helps correcting the grammatical errors and rephrasing the content, if needed.")
st.write("___________________________")

# Form 
uploaded_file = st.file_uploader("Choose a doc file to upload", type=['docx'])
action = st.radio("Select the action you want to take", ('Correction', 'Translation'))

if action == "Correction":
    query = "Please correct the grammar of the article in the same language and make it better"
elif action == "Translation":
    lang = st.radio("Please select the translation language of your document?", ("Hindi", "English", "Urdu"))
    if lang:
        query = "Please translation the article to " + lang + " and make it better also separate paragraphs"
elif action == 'Custom':
    query = st.text_input("Write the prompt")

thread_status = st.checkbox('Do you want to generate an automated thread of the report?')

# Button to search
if st.button("Analyze"):
    if query != "":
        if uploaded_file is not None:
            load_screen = 1
            while load_screen == 1:
                with st.spinner("Processing data, please wait!"):
                    doc_text, heading = load_doc(uploaded_file)
                    # Count the number of words
                    word_count = len(doc_text.split())
                    # Check if the word count exceeds 500
                    if word_count > 500:
                        st.write("Total words in the file: " + str(word_count))
                        st.error("Error: Document word count exceeds more than 500 words")
                        raise SystemExit
                    input = doc_text + " " + query
                    answer = chatgpt_api(input)
                    load_screen = 0
            st.write("______________________________________")
            st.subheader(action)
            st.write(answer)
            all_docs = []
            if thread_status:
                load_screen = 1
                while load_screen == 1:
                    with st.spinner("Generating thread please wait!"):
                        thread_answer = gen_thread(answer)
                        load_screen = 0
                        st.write("______________________________________")
                        st.subheader("Thread of the Article")
                        st.write(thread_answer)
                thread_doc = save_doc(thread_answer, "Thread")
                all_docs.append(thread_doc)
            article_doc = save_doc(answer, action)
            all_docs.append(article_doc)
            save_multiple_doc(all_docs)
        else:
            st.error("Please upload a document to perform any action")
else:
    pass